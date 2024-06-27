#  Licensed under the Apache License, Version 2.0 (the "License");
#  you may not use this file except in compliance with the License.
#  You may obtain a copy of the License at
#
#      http://www.apache.org/licenses/LICENSE-2.0
#
#  Unless required by applicable law or agreed to in writing, software
#  distributed under the License is distributed on an "AS IS" BASIS,
#  WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
#  See the License for the specific language governing permissions and
#  limitations under the License.
#
import copy
from tika import parser
import re
from io import BytesIO
from docx import Document

from api.db import ParserType
from rag.nlp import bullets_category, is_english, tokenize, remove_contents_table, hierarchical_merge, \
    make_colon_as_title, add_positions, tokenize_chunks, find_codec, docx_question_level
from rag.nlp import rag_tokenizer
from deepdoc.parser import PdfParser, DocxParser, PlainParser, HtmlParser
from rag.settings import cron_logger


class Docx(DocxParser):
    def __init__(self):
        pass

    def __clean(self, line):
        line = re.sub(r"\u3000", " ", line).strip()
        return line

    def old_call(self, filename, binary=None, from_page=0, to_page=100000):
        self.doc = Document(
            filename) if not binary else Document(BytesIO(binary))
        pn = 0
        lines = []
        for p in self.doc.paragraphs:
            if pn > to_page:
                break
            if from_page <= pn < to_page and p.text.strip():
                lines.append(self.__clean(p.text))
            for run in p.runs:
                if 'lastRenderedPageBreak' in run._element.xml:
                    pn += 1
                    continue
                if 'w:br' in run._element.xml and 'type="page"' in run._element.xml:
                    pn += 1
        return [l for l in lines if l]

    def __call__(self, filename, binary=None, from_page=0, to_page=100000):
        self.doc = Document(
            filename) if not binary else Document(BytesIO(binary))
        pn = 0
        last_question, last_answer, last_level = "", "", -1
        lines = []
        root = DocxNode()
        point = root
        bull = bullets_category([p.text for p in self.doc.paragraphs])
        for p in self.doc.paragraphs:
            if pn > to_page:
                break
            question_level, p_text = 0, ''
            if from_page <= pn < to_page and p.text.strip():
                question_level, p_text = docx_question_level(p, bull)
            if not question_level or question_level > 6: # not a question
                last_answer = f'{last_answer}\n{p_text}'
            else:   # is a question
                if last_question:
                    while last_level <= point.level:
                        point = point.parent
                    new_node = DocxNode(last_question, last_answer, last_level, [], point)
                    point.childs.append(new_node)
                    point = new_node
                    last_question, last_answer, last_level = '', '', -1
                last_level = question_level
                last_question = p_text
            
            for run in p.runs:
                if 'lastRenderedPageBreak' in run._element.xml:
                    pn += 1
                    continue
                if 'w:br' in run._element.xml and 'type="page"' in run._element.xml:
                    pn += 1
        if last_question:
            while last_level <= point.level:
                point = point.parent
            new_node = DocxNode(last_question, last_answer, last_level, [], point)
            point.childs.append(new_node)
            point = new_node
            last_question, last_answer, last_level = '', '', -1
        traversal_queue = [root]
        while traversal_queue:
            current_node: DocxNode = traversal_queue.pop()
            sum_text = f'{self.__clean(current_node.question)}\n{self.__clean(current_node.answer)}'
            if not current_node.childs:
                continue
            for child in current_node.childs:
                sum_text = f'{sum_text}\n{self.__clean(child.question)}'
                traversal_queue.insert(0, child)
            lines.append(self.__clean(sum_text))
        return [l for l in lines if l]
class DocxNode:
    def __init__(self, question: str = '', answer: str = '', level: int = 0, childs: list = [], parent = None) -> None:
        self.question = question
        self.answer = answer
        self.level = level
        self.childs = childs
        self.parent = parent
    def __str__(self) -> str:
        return f'''
            question:{self.question},
            answer:{self.answer},
            level:{self.level},
            childs:{self.childs}
        '''


class Pdf(PdfParser):
    def __init__(self):
        self.model_speciess = ParserType.LAWS.value
        super().__init__()

    def __call__(self, filename, binary=None, from_page=0,
                 to_page=100000, zoomin=3, callback=None):
        callback(msg="OCR is running...")
        self.__images__(
            filename if not binary else binary,
            zoomin,
            from_page,
            to_page,
            callback
        )
        callback(msg="OCR finished")

        from timeit import default_timer as timer
        start = timer()
        self._layouts_rec(zoomin)
        callback(0.67, "Layout analysis finished")
        cron_logger.info("layouts:".format(
            (timer() - start) / (self.total_page + 0.1)))
        self._naive_vertical_merge()

        callback(0.8, "Text extraction finished")

        return [(b["text"], self._line_tag(b, zoomin))
                for b in self.boxes], None


def chunk(filename, binary=None, from_page=0, to_page=100000,
          lang="Chinese", callback=None, **kwargs):
    """
        Supported file formats are docx, pdf, txt.
    """
    doc = {
        "docnm_kwd": filename,
        "title_tks": rag_tokenizer.tokenize(re.sub(r"\.[a-zA-Z]+$", "", filename))
    }
    doc["title_sm_tks"] = rag_tokenizer.fine_grained_tokenize(doc["title_tks"])
    pdf_parser = None
    sections = []
    # is it English
    eng = lang.lower() == "english"  # is_english(sections)

    if re.search(r"\.docx$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        for txt in Docx()(filename, binary):
            sections.append(txt)
        callback(0.8, "Finish parsing.")
        chunks = sections
        return tokenize_chunks(chunks, doc, eng, pdf_parser)

    elif re.search(r"\.pdf$", filename, re.IGNORECASE):
        pdf_parser = Pdf() if kwargs.get(
            "parser_config", {}).get(
            "layout_recognize", True) else PlainParser()
        for txt, poss in pdf_parser(filename if not binary else binary,
                                    from_page=from_page, to_page=to_page, callback=callback)[0]:
            sections.append(txt + poss)

    elif re.search(r"\.txt$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        txt = ""
        if binary:
            encoding = find_codec(binary)
            txt = binary.decode(encoding, errors="ignore")
        else:
            with open(filename, "r") as f:
                while True:
                    l = f.readline()
                    if not l:
                        break
                    txt += l
        sections = txt.split("\n")
        sections = [l for l in sections if l]
        callback(0.8, "Finish parsing.")

    elif re.search(r"\.(htm|html)$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        sections = HtmlParser()(filename, binary)
        sections = [l for l in sections if l]
        callback(0.8, "Finish parsing.")

    elif re.search(r"\.doc$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        binary = BytesIO(binary)
        doc_parsed = parser.from_buffer(binary)
        sections = doc_parsed['content'].split('\n')
        sections = [l for l in sections if l]
        callback(0.8, "Finish parsing.")

    else:
        raise NotImplementedError(
            "file type not supported yet(doc, docx, pdf, txt supported)")


    # Remove 'Contents' part
    remove_contents_table(sections, eng)

    make_colon_as_title(sections)
    bull = bullets_category(sections)
    chunks = hierarchical_merge(bull, sections, 5)
    if not chunks:
        callback(0.99, "No chunk parsed out.")

    return tokenize_chunks(["\n".join(ck)
                           for ck in chunks], doc, eng, pdf_parser)


if __name__ == "__main__":
    import sys

    def dummy(prog=None, msg=""):
        pass
    chunk(sys.argv[1], callback=dummy)
